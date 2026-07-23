from collections import Counter, defaultdict
from datetime import datetime
from io import BytesIO
import json
import os
from pathlib import Path
from threading import Lock

from flask import Flask, render_template, jsonify, request, send_file
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.shared import Inches, Pt, RGBColor

app = Flask(__name__)


STATE_LOCK = Lock()
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = Path(os.getenv('APP_DATA_DIR') or os.getenv('RENDER_DISK_ROOT') or (BASE_DIR / 'data'))
STATE_FILE = DATA_DIR / 'app_state.json'


def load_state():
    if not STATE_FILE.exists():
        return {
            'diligencias': [],
            'processos': [],
            'report_history': [],
        }

    try:
        data = json.loads(STATE_FILE.read_text(encoding='utf-8'))
    except (json.JSONDecodeError, OSError):
        return {
            'diligencias': [],
            'processos': [],
            'report_history': [],
        }

    return {
        'diligencias': data.get('diligencias', []) if isinstance(data.get('diligencias', []), list) else [],
        'processos': data.get('processos', []) if isinstance(data.get('processos', []), list) else [],
        'report_history': data.get('report_history', []) if isinstance(data.get('report_history', []), list) else [],
    }


def save_state():
    payload = {
        'diligencias': diligencias,
        'processos': processos,
        'report_history': report_history,
    }

    with STATE_LOCK:
        DATA_DIR.mkdir(parents=True, exist_ok=True)
        tmp_file = STATE_FILE.with_suffix('.tmp')
        tmp_file.write_text(json.dumps(payload, ensure_ascii=False), encoding='utf-8')
        tmp_file.replace(STATE_FILE)


def set_cell_background(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def format_currency_brl(value):
    if value is None:
        return '-'
    text = f'{value:,.2f}'
    return f'R$ {text}'.replace(',', 'X').replace('.', ',').replace('X', '.')


def format_distance_km(value):
    if value is None:
        return '-'

    text = f'{value:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')
    return f'{text} km'


def build_operational_report_text(item):
    details = []

    if item.get('valor_causa') is not None:
        details.append(f"Valor da causa: {format_currency_brl(item.get('valor_causa'))}")
    if item.get('modalidade_diligencia'):
        details.append(f"Modalidade: {item.get('modalidade_diligencia')}")
    if item.get('distancia_roteiro') is not None:
        details.append(f"Distância: {format_distance_km(item.get('distancia_roteiro'))}")
    if item.get('preco_gasolina') is not None:
        details.append(f"Gasolina: {format_currency_brl(item.get('preco_gasolina'))}")
    if item.get('preco_aluguel_carro') is not None:
        details.append(f"Aluguel: {format_currency_brl(item.get('preco_aluguel_carro'))}")
    if item.get('roteiro_estrategico'):
        details.append(f"Roteiro: {item.get('roteiro_estrategico')}")
    if item.get('modus_operandi'):
        details.append(f"Operação: {item.get('modus_operandi')}")

    return ' | '.join(details) if details else '-'


def build_docx_report(data):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('RELATORIO GERENCIAL DE DILIGENCIAS')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(31, 41, 55)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('Diligencia Elite RJ')
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(55, 65, 81)

    period = document.add_paragraph()
    period.alignment = WD_ALIGN_PARAGRAPH.CENTER
    period_run = period.add_run(datetime.now().strftime('Emitido em %d/%m/%Y as %H:%M'))
    period_run.italic = True
    period_run.font.size = Pt(10)
    period_run.font.color.rgb = RGBColor(107, 114, 128)

    document.add_paragraph()
    summary_table = document.add_table(rows=2, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary_table.style = 'Table Grid'

    total = len(data)
    status_counter = Counter([item.get('status', 'Nao Informado') for item in data])
    municipios = len(set(item.get('municipio', '') for item in data if item.get('municipio')))
    total_alvara = sum(item.get('valor_alvara') or 0 for item in data)

    summary_headers = ['Total de Processos', 'Municipios Ativos', 'Urgentes', 'Total Alvaras']
    summary_values = [
        str(total),
        str(municipios),
        str(status_counter.get('Urgente', 0)),
        format_currency_brl(total_alvara),
    ]

    for idx, value in enumerate(summary_headers):
        cell = summary_table.cell(0, idx)
        cell.text = value
        set_cell_background(cell, 'E5E7EB')
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for idx, value in enumerate(summary_values):
        cell = summary_table.cell(1, idx)
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(30, 64, 175)

    document.add_page_break()

    heading = document.add_heading('1. Controle Operacional das Diligencias', level=1)
    heading.runs[0].font.color.rgb = RGBColor(31, 41, 55)

    table = document.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['#', 'Processo', 'Responsavel', 'Status', 'Regiao', 'Municipio', 'Comarca', 'Valor Alvara', 'Resumo']

    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_background(cell, 'DBEAFE')
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for idx, item in enumerate(data, start=1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = item.get('numero', '')
        row[2].text = item.get('responsavel', '')
        row[3].text = item.get('status', '')
        row[4].text = item.get('region', '')
        row[5].text = item.get('municipio', '')
        row[6].text = item.get('comarca', '')
        row[7].text = format_currency_brl(item.get('valor_alvara'))
        row[8].text = ' | '.join(filter(None, [item.get('resumo', ''), build_operational_report_text(item)]))

    if data:
        document.add_paragraph()
        top_heading = document.add_heading('2. Separacao por Comarca', level=2)
        top_heading.runs[0].font.color.rgb = RGBColor(31, 41, 55)
        comarca_groups = defaultdict(list)
        for item in data:
            comarca_name = item.get('comarca', '').strip() or 'Sem comarca'
            comarca_groups[comarca_name].append(item)

        for comarca_name in sorted(comarca_groups.keys()):
            section_heading = document.add_heading(f'Comarca: {comarca_name}', level=3)
            section_heading.runs[0].font.color.rgb = RGBColor(55, 65, 81)
            group_table = document.add_table(rows=1, cols=6)
            group_table.style = 'Table Grid'
            group_headers = ['Processo', 'Responsavel', 'Status', 'Municipio', 'Valor Alvara', 'Resumo']

            for idx, header in enumerate(group_headers):
                cell = group_table.rows[0].cells[idx]
                cell.text = header
                set_cell_background(cell, 'E5E7EB')
                for run in cell.paragraphs[0].runs:
                    run.bold = True

            for item in comarca_groups[comarca_name]:
                row = group_table.add_row().cells
                row[0].text = item.get('numero', '')
                row[1].text = item.get('responsavel', '')
                row[2].text = item.get('status', '')
                row[3].text = item.get('municipio', '')
                row[4].text = format_currency_brl(item.get('valor_alvara'))
                row[5].text = ' | '.join(filter(None, [item.get('resumo', ''), build_operational_report_text(item)]))

            document.add_paragraph()

        top_municipios_heading = document.add_heading('3. Top Municipios por Volume', level=2)
        top_municipios_heading.runs[0].font.color.rgb = RGBColor(31, 41, 55)
        top_items = Counter(item.get('municipio', 'Nao informado') for item in data).most_common(5)
        top_table = document.add_table(rows=1, cols=2)
        top_table.style = 'Table Grid'
        top_table.rows[0].cells[0].text = 'Municipio'
        top_table.rows[0].cells[1].text = 'Quantidade'
        set_cell_background(top_table.rows[0].cells[0], 'E5E7EB')
        set_cell_background(top_table.rows[0].cells[1], 'E5E7EB')
        for hcell in top_table.rows[0].cells:
            for run in hcell.paragraphs[0].runs:
                run.bold = True

        for municipio, qtd in top_items:
            cells = top_table.add_row().cells
            cells[0].text = municipio
            cells[1].text = str(qtd)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run('Diligencia Elite RJ | Relatorio automatico')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(107, 114, 128)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

municipio_coordinates = {
    'Angra dos Reis': (-23.0066, -44.3182),
    'Aperibé': (-21.4219, -42.0945),
    'Araruama': (-22.8750, -42.3450),
    'Areal': (-22.4583, -43.3667),
    'Armação dos Búzios': (-22.7545, -41.8863),
    'Arraial do Cabo': (-22.9656, -42.0281),
    'Barra do Piraí': (-22.4781, -43.8219),
    'Barra Mansa': (-22.5447, -44.1780),
    'Belford Roxo': (-22.7645, -43.3998),
    'Bom Jardim': (-22.3081, -42.1781),
    'Bom Jesus do Itabapoana': (-21.1367, -41.7769),
    'Cabo Frio': (-22.8833, -42.0333),
    'Cachoeiras de Macacu': (-22.4667, -42.6333),
    'Cambuci': (-21.5203, -41.8964),
    'Campos dos Goytacazes': (-21.7618, -41.3280),
    'Cantagalo': (-21.9583, -42.4283),
    'Carapebus': (-21.3428, -41.2267),
    'Cardoso Moreira': (-21.0247, -41.5889),
    'Carmo': (-21.6281, -42.0542),
    'Casimiro de Abreu': (-22.4333, -42.0950),
    'Comendador Levy Gasparian': (-21.8131, -42.3642),
    'Conceição de Macabu': (-21.3333, -41.6167),
    'Cordeiro': (-22.1333, -42.2667),
    'Duas Barras': (-22.0667, -42.0933),
    'Duque de Caxias': (-22.7858, -43.3040),
    'Engenheiro Paulo de Frontin': (-22.5639, -43.6486),
    'Guapimirim': (-22.5056, -43.0394),
    'Iguaba Grande': (-22.8194, -42.3306),
    'Itaboraí': (-22.7447, -42.8606),
    'Itaguaí': (-22.8556, -43.7722),
    'Italva': (-21.1556, -41.8347),
    'Itaocara': (-21.2447, -42.0639),
    'Itaperuna': (-21.2264, -41.9033),
    'Itatiaia': (-22.4867, -44.5533),
    'Japeri': (-22.6622, -43.7319),
    'Laje do Muriaé': (-21.2719, -42.1722),
    'Macaé': (-22.3817, -41.5919),
    'Macuco': (-21.3906, -41.8142),
    'Magé': (-22.6496, -43.0350),
    'Mangaratiba': (-23.1667, -44.0167),
    'Maricá': (-22.9658, -42.8245),
    'Mendes': (-22.5892, -43.7511),
    'Mesquita': (-22.7703, -43.5572),
    'Miguel Pereira': (-22.5167, -43.5167),
    'Miracema': (-21.4667, -42.1167),
    'Natividade': (-21.3222, -41.9778),
    'Nilópolis': (-22.8061, -43.4263),
    'Niterói': (-22.8832, -43.1034),
    'Nova Friburgo': (-22.2833, -42.5333),
    'Nova Iguaçu': (-22.7592, -43.4514),
    'Paracambi': (-22.6083, -43.7056),
    'Paraíba do Sul': (-21.2161, -41.5942),
    'Paraty': (-23.2180, -44.4194),
    'Paty do Alferes': (-22.3708, -43.6389),
    'Petrópolis': (-22.5244, -43.1789),
    'Pinheiral': (-22.5244, -43.9972),
    'Piraí': (-22.6417, -44.0250),
    'Porciúncula': (-21.5161, -42.0500),
    'Porto Real': (-22.5250, -44.2944),
    'Quatis': (-22.5233, -44.1867),
    'Queimados': (-22.7417, -43.6139),
    'Quissamã': (-21.8656, -41.0614),
    'Resende': (-22.4717, -44.4500),
    'Rio Bonito': (-22.7122, -42.5544),
    'Rio Claro': (-22.2717, -44.0133),
    'Rio das Flores': (-22.1833, -43.8833),
    'Rio das Ostras': (-22.5286, -41.7667),
    'Rio de Janeiro': (-22.9068, -43.1729),
    'Santa Maria Madalena': (-21.9600, -42.0367),
    'Santo Antônio de Pádua': (-21.5419, -41.9569),
    'São Fidélis': (-21.6400, -41.6400),
    'São Francisco de Itabapoana': (-21.3292, -41.0892),
    'São Gonçalo': (-22.8268, -43.0532),
    'São João da Barra': (-21.6419, -41.0539),
    'São João de Meriti': (-22.8183, -43.3683),
    'São José de Ubá': (-21.1306, -41.5153),
    'São José do Vale do Rio Preto': (-22.3772, -43.0772),
    'São Pedro da Aldeia': (-22.8333, -42.0667),
    'São Sebastião do Alto': (-22.0833, -42.2333),
    'Sapucaia': (-22.3833, -42.8500),
    'Saquarema': (-22.9211, -42.5000),
    'Seropédica': (-22.7383, -43.8047),
    'Silva Jardim': (-22.7583, -42.7744),
    'Sumidouro': (-22.0833, -42.3333),
    'Tanguá': (-22.6756, -42.7219),
    'Teresópolis': (-22.4139, -43.0711),
    'Trajano de Moraes': (-21.9333, -42.1333),
    'Três Rios': (-22.2167, -43.2167),
    'Valença': (-22.3083, -43.7056),
    'Varre-Sai': (-21.0292, -42.0167),
    'Vassouras': (-22.3844, -43.6644),
    'Volta Redonda': (-22.5139, -44.0967),
}

def get_municipio_coords(name):
    return municipio_coordinates.get(name, (-22.9068, -43.1729))


def parse_optional_money(value):
    if value in (None, ''):
        return None

    try:
        amount = float(value)
    except (TypeError, ValueError):
        return None

    if amount < 0:
        return None

    return round(amount, 2)


def parse_optional_text(value, fallback=''):
    if value is None:
        return fallback

    text = str(value).strip()
    return text if text else fallback


def parse_optional_distance(value):
    if value in (None, ''):
        return None

    try:
        distance = float(value)
    except (TypeError, ValueError):
        return None

    if distance < 0:
        return None

    return round(distance, 2)


def extract_operational_fields(data, fallback=None):
    fallback = fallback or {}
    return {
        'valor_causa': parse_optional_money(data.get('valor_causa', fallback.get('valor_causa'))),
        'roteiro_estrategico': parse_optional_text(data.get('roteiro_estrategico'), fallback.get('roteiro_estrategico', '')),
        'modalidade_diligencia': parse_optional_text(data.get('modalidade_diligencia'), fallback.get('modalidade_diligencia', 'Não informado')),
        'distancia_roteiro': parse_optional_distance(data.get('distancia_roteiro', fallback.get('distancia_roteiro'))),
        'preco_gasolina': parse_optional_money(data.get('preco_gasolina', fallback.get('preco_gasolina'))),
        'preco_aluguel_carro': parse_optional_money(data.get('preco_aluguel_carro', fallback.get('preco_aluguel_carro'))),
        'modus_operandi': parse_optional_text(data.get('modus_operandi'), fallback.get('modus_operandi', '')),
    }


def build_total_value(valor, despesas):
    amounts = [amount for amount in (valor, despesas) if amount is not None]
    if not amounts:
        return None
    return round(sum(amounts), 2)


def parse_optional_float(value, fallback):
    try:
        return float(value)
    except (TypeError, ValueError):
        return fallback


def parse_positive_int(value, fallback=1):
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        return fallback
    return parsed if parsed > 0 else fallback


def normalize_state(state):
    stored_diligencias = state.get('diligencias', []) if isinstance(state.get('diligencias', []), list) else []
    stored_processos = state.get('processos', []) if isinstance(state.get('processos', []), list) else []

    # Deduplicate by process number and rebuild cross-links to keep table/map consistent.
    processos_by_numero = {}
    for item in stored_processos:
        numero = str(item.get('numero', '')).strip()
        if not numero:
            continue
        processos_by_numero[numero] = {
            'id': item.get('id'),
            'numero': numero,
            'status': item.get('status', 'Pendente'),
            'region': item.get('region', 'Metropolitana'),
            'municipio': item.get('municipio', ''),
            'comarca': item.get('comarca', ''),
            'responsavel': item.get('responsavel', ''),
            'urgencia': item.get('urgencia', item.get('status', 'Pendente')),
            'resumo': item.get('resumo', ''),
            'valor_alvara': parse_optional_money(item.get('valor_alvara', item.get('valor'))),
            'valor_total': parse_optional_money(item.get('valor_total', item.get('valor_alvara', item.get('valor')))),
            'valor_causa': parse_optional_money(item.get('valor_causa')),
            'roteiro_estrategico': item.get('roteiro_estrategico', ''),
            'modalidade_diligencia': item.get('modalidade_diligencia', 'Não informado'),
            'distancia_roteiro': parse_optional_distance(item.get('distancia_roteiro')),
            'preco_gasolina': parse_optional_money(item.get('preco_gasolina')),
            'preco_aluguel_carro': parse_optional_money(item.get('preco_aluguel_carro')),
            'modus_operandi': item.get('modus_operandi', ''),
        }

    diligencias_by_numero = {}
    for item in stored_diligencias:
        numero = str(item.get('process_number', '')).strip()
        if not numero:
            continue
        lat, lng = get_municipio_coords(item.get('municipio', ''))
        diligencias_by_numero[numero] = {
            'id': item.get('id'),
            'name': item.get('name', '').strip() or 'Nova diligência',
            'process_number': numero,
            'responsavel': item.get('responsavel', '').strip(),
            'region': item.get('region', 'Não informado'),
            'municipio': item.get('municipio', ''),
            'comarca': item.get('comarca', ''),
            'lat': parse_optional_float(item.get('lat'), lat),
            'lng': parse_optional_float(item.get('lng'), lng),
            'status': item.get('status', 'Pendente'),
            'resumo': item.get('resumo', ''),
            'processos': parse_positive_int(item.get('processos', 1)),
            'valor_alvara': parse_optional_money(item.get('valor_alvara', item.get('valor'))),
            'valor_total': parse_optional_money(item.get('valor_total', item.get('valor_alvara', item.get('valor')))),
            'valor_causa': parse_optional_money(item.get('valor_causa')),
            'roteiro_estrategico': item.get('roteiro_estrategico', ''),
            'modalidade_diligencia': item.get('modalidade_diligencia', 'Não informado'),
            'distancia_roteiro': parse_optional_distance(item.get('distancia_roteiro')),
            'preco_gasolina': parse_optional_money(item.get('preco_gasolina')),
            'preco_aluguel_carro': parse_optional_money(item.get('preco_aluguel_carro')),
            'modus_operandi': item.get('modus_operandi', ''),
        }

    for numero, proc in processos_by_numero.items():
        if numero not in diligencias_by_numero:
            lat, lng = get_municipio_coords(proc.get('municipio', ''))
            diligencias_by_numero[numero] = {
                'id': None,
                'name': proc.get('comarca') or proc.get('numero') or 'Nova diligência',
                'process_number': numero,
                'responsavel': proc.get('responsavel', ''),
                'region': proc.get('region', 'Não informado'),
                'municipio': proc.get('municipio', ''),
                'comarca': proc.get('comarca', ''),
                'lat': lat,
                'lng': lng,
                'status': proc.get('status', 'Pendente'),
                'resumo': proc.get('resumo', ''),
                'processos': 1,
                'valor_alvara': proc.get('valor_alvara'),
                'valor_total': proc.get('valor_total'),
                'valor_causa': proc.get('valor_causa'),
                'roteiro_estrategico': proc.get('roteiro_estrategico', ''),
                'modalidade_diligencia': proc.get('modalidade_diligencia', 'Não informado'),
                'distancia_roteiro': proc.get('distancia_roteiro'),
                'preco_gasolina': proc.get('preco_gasolina'),
                'preco_aluguel_carro': proc.get('preco_aluguel_carro'),
                'modus_operandi': proc.get('modus_operandi', ''),
            }

    for numero, dil in diligencias_by_numero.items():
        if numero not in processos_by_numero:
            processos_by_numero[numero] = {
                'id': None,
                'numero': numero,
                'status': dil.get('status', 'Pendente'),
                'region': dil.get('region', 'Metropolitana'),
                'municipio': dil.get('municipio', ''),
                'comarca': dil.get('comarca', ''),
                'responsavel': dil.get('responsavel', ''),
                'urgencia': dil.get('status', 'Pendente'),
                'resumo': dil.get('resumo', ''),
                'valor_alvara': dil.get('valor_alvara'),
                'valor_total': dil.get('valor_total'),
                'valor_causa': dil.get('valor_causa'),
                'roteiro_estrategico': dil.get('roteiro_estrategico', ''),
                'modalidade_diligencia': dil.get('modalidade_diligencia', 'Não informado'),
                'distancia_roteiro': dil.get('distancia_roteiro'),
                'preco_gasolina': dil.get('preco_gasolina'),
                'preco_aluguel_carro': dil.get('preco_aluguel_carro'),
                'modus_operandi': dil.get('modus_operandi', ''),
            }

    normalized_processos = []
    next_id = 1
    for _, item in sorted(processos_by_numero.items(), key=lambda entry: entry[0]):
        normalized = dict(item)
        normalized['id'] = next_id
        normalized_processos.append(normalized)
        next_id += 1

    normalized_diligencias = []
    next_id = 1
    for _, item in sorted(diligencias_by_numero.items(), key=lambda entry: entry[0]):
        normalized = dict(item)
        normalized['id'] = next_id
        normalized_diligencias.append(normalized)
        next_id += 1

    normalized_history = state.get('report_history', []) if isinstance(state.get('report_history', []), list) else []
    return {
        'diligencias': normalized_diligencias,
        'processos': normalized_processos,
        'report_history': normalized_history,
    }

_initial_state = normalize_state(load_state())
diligencias = _initial_state['diligencias']
processos = _initial_state['processos']
report_history = _initial_state['report_history']

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/health')
def health_check():
    return jsonify({'status': 'ok'})

@app.route('/api/municipios-coords', methods=['GET'])
def get_municipios_coords():
    return jsonify({name: {'lat': coords[0], 'lng': coords[1]} for name, coords in municipio_coordinates.items()})

@app.route('/api/diligencias', methods=['GET', 'POST'])
def get_diligencias():
    if request.method == 'POST':
        data = request.get_json() or {}
        process_number = data.get('process_number', '').strip()
        if not process_number:
            return jsonify({'error': 'Número do processo é obrigatório'}), 400

        existing = next((d for d in diligencias if d.get('process_number') == process_number), None)
        valor_alvara = parse_optional_money(data.get('valor_alvara', data.get('valor')))
        operational_fields = extract_operational_fields(data, existing)
        if existing:
            lat, lng = get_municipio_coords(data.get('municipio', existing.get('municipio', '')))
            existing.update({
                'name': data.get('name', existing.get('name', '')).strip() or 'Nova diligência',
                'responsavel': data.get('responsavel', existing.get('responsavel', '')).strip(),
                'region': data.get('region', existing.get('region', 'Não informado')),
                'municipio': data.get('municipio', existing.get('municipio', '')),
                'comarca': data.get('comarca', existing.get('comarca', '')),
                'lat': parse_optional_float(data.get('lat'), lat),
                'lng': parse_optional_float(data.get('lng'), lng),
                'status': data.get('status', existing.get('status', 'Pendente')),
                'resumo': data.get('summary', existing.get('resumo', '')).strip(),
                'processos': parse_positive_int(data.get('processos', existing.get('processos', 1))),
                'valor_alvara': valor_alvara,
                'valor_total': valor_alvara,
            })
            existing.update(operational_fields)
            item = existing
        else:
            item = {
                'id': max((d['id'] for d in diligencias), default=0) + 1,
                'name': data.get('name', '').strip() or 'Nova diligência',
                'process_number': process_number,
                'responsavel': data.get('responsavel', '').strip(),
                'region': data.get('region', 'Não informado'),
                'municipio': data.get('municipio', ''),
                'comarca': data.get('comarca', ''),
                'lat': parse_optional_float(data.get('lat'), -22.9068),
                'lng': parse_optional_float(data.get('lng'), -43.1729),
                'status': data.get('status', 'Pendente'),
                'resumo': data.get('summary', '').strip(),
                'processos': parse_positive_int(data.get('processos', 1)),
                'valor_alvara': valor_alvara,
                'valor_total': valor_alvara,
                **operational_fields,
            }
            diligencias.append(item)

        process_to_update = next((p for p in processos if p.get('numero') == process_number), None)
        if process_to_update:
            process_to_update.update({
                'status': item.get('status', process_to_update.get('status', 'Pendente')),
                'region': item.get('region', process_to_update.get('region', 'Metropolitana')),
                'municipio': item.get('municipio', process_to_update.get('municipio', '')),
                'comarca': item.get('comarca', process_to_update.get('comarca', '')),
                'responsavel': item.get('responsavel', process_to_update.get('responsavel', '')),
                'urgencia': item.get('status', process_to_update.get('urgencia', 'Pendente')),
                'resumo': item.get('resumo', process_to_update.get('resumo', '')),
                'valor_alvara': item.get('valor_alvara'),
                'valor_total': item.get('valor_total'),
            })
            process_to_update.update(extract_operational_fields(item, process_to_update))
        else:
            processos.append({
                'id': max((p['id'] for p in processos), default=0) + 1,
                'numero': process_number,
                'status': item.get('status', 'Pendente'),
                'region': item.get('region', 'Metropolitana'),
                'municipio': item.get('municipio', ''),
                'comarca': item.get('comarca', ''),
                'responsavel': item.get('responsavel', ''),
                'urgencia': item.get('status', 'Pendente'),
                'resumo': item.get('resumo', ''),
                'valor_alvara': item.get('valor_alvara'),
                'valor_total': item.get('valor_total'),
                **extract_operational_fields(item),
            })

        save_state()
        return jsonify(item), 201

    return jsonify(diligencias)

@app.route('/api/processos', methods=['GET', 'POST'])
def get_processos():
    if request.method == 'POST':
        data = request.get_json() or {}
        numero = data.get('process_number', '').strip()
        if not numero:
            return jsonify({'error': 'Número do processo é obrigatório'}), 400

        existing = next((p for p in processos if p.get('numero') == numero), None)
        valor_alvara = parse_optional_money(data.get('valor_alvara', data.get('valor')))
        operational_fields = extract_operational_fields(data, existing)
        if existing:
            existing.update({
                'status': data.get('status', existing.get('status', 'Pendente')),
                'region': data.get('region', existing.get('region', 'Metropolitana')),
                'municipio': data.get('municipio', existing.get('municipio', '')),
                'comarca': data.get('comarca', existing.get('comarca', '')),
                'responsavel': data.get('responsavel', existing.get('responsavel', '')),
                'urgencia': data.get('status', existing.get('urgencia', 'Pendente')),
                'resumo': data.get('summary', existing.get('resumo', '')).strip(),
                'valor_alvara': valor_alvara,
                'valor_total': valor_alvara,
            })
            existing.update(operational_fields)
            item = existing
        else:
            item = {
                'id': max((p['id'] for p in processos), default=0) + 1,
                'numero': numero,
                'status': data.get('status', 'Pendente'),
                'region': data.get('region', 'Metropolitana'),
                'municipio': data.get('municipio', ''),
                'comarca': data.get('comarca', ''),
                'responsavel': data.get('responsavel', ''),
                'urgencia': data.get('status', 'Pendente'),
                'resumo': data.get('summary', '').strip(),
                'valor_alvara': valor_alvara,
                'valor_total': valor_alvara,
                **operational_fields,
            }
            processos.append(item)

        matching_diligencia = next((d for d in diligencias if d.get('process_number') == numero), None)
        lat, lng = get_municipio_coords(item.get('municipio', ''))
        if matching_diligencia:
            matching_diligencia.update({
                'name': matching_diligencia.get('name') or item.get('comarca') or item.get('numero') or 'Nova diligência',
                'responsavel': item.get('responsavel', ''),
                'region': item.get('region', 'Não informado'),
                'municipio': item.get('municipio', ''),
                'comarca': item.get('comarca', ''),
                'lat': parse_optional_float(matching_diligencia.get('lat'), lat),
                'lng': parse_optional_float(matching_diligencia.get('lng'), lng),
                'status': item.get('status', 'Pendente'),
                'resumo': item.get('resumo', ''),
                'valor_alvara': item.get('valor_alvara'),
                'valor_total': item.get('valor_total'),
            })
            matching_diligencia.update(extract_operational_fields(item, matching_diligencia))
        else:
            diligencias.append({
                'id': max((d['id'] for d in diligencias), default=0) + 1,
                'name': item.get('comarca') or item.get('numero') or 'Nova diligência',
                'process_number': numero,
                'responsavel': item.get('responsavel', ''),
                'region': item.get('region', 'Não informado'),
                'municipio': item.get('municipio', ''),
                'comarca': item.get('comarca', ''),
                'lat': lat,
                'lng': lng,
                'status': item.get('status', 'Pendente'),
                'resumo': item.get('resumo', ''),
                'processos': 1,
                'valor_alvara': item.get('valor_alvara'),
                'valor_total': item.get('valor_total'),
                **extract_operational_fields(item),
            })

        save_state()
        return jsonify(item), 201

    return jsonify(processos)

@app.route('/api/processos/<int:process_id>', methods=['PUT'])
def update_processo(process_id):
    data = request.get_json() or {}
    process_to_update = next((p for p in processos if p['id'] == process_id), None)
    if process_to_update is None:
        return jsonify({'error': 'Processo não encontrado'}), 404

    old_numero = process_to_update['numero']
    valor_alvara = parse_optional_money(data.get('valor_alvara', data.get('valor', process_to_update.get('valor_alvara'))))
    operational_fields = extract_operational_fields(data, process_to_update)
    process_to_update.update({
        'numero': data.get('process_number', process_to_update['numero']).strip(),
        'status': data.get('status', process_to_update['status']),
        'region': data.get('region', process_to_update.get('region', 'Metropolitana')),
        'municipio': data.get('municipio', process_to_update['municipio']),
        'comarca': data.get('comarca', process_to_update['comarca']),
        'responsavel': data.get('responsavel', process_to_update['responsavel']),
        'urgencia': data.get('status', process_to_update['urgencia']),
        'resumo': data.get('summary', process_to_update['resumo']).strip(),
        'valor_alvara': valor_alvara,
        'valor_total': valor_alvara,
    })
    process_to_update.update(operational_fields)

    updated_numero = process_to_update['numero']
    matching_diligencia = next((d for d in diligencias if d.get('process_number') == old_numero), None)
    if matching_diligencia:
        matching_diligencia.update({
            'process_number': process_to_update['numero'],
            'responsavel': process_to_update['responsavel'],
            'region': data.get('region', matching_diligencia.get('region')),
            'municipio': process_to_update['municipio'],
            'comarca': process_to_update['comarca'],
            'status': process_to_update['status'],
            'resumo': process_to_update['resumo'],
            'valor_alvara': process_to_update.get('valor_alvara'),
            'valor_total': process_to_update.get('valor_total'),
        })
        matching_diligencia.update(extract_operational_fields(process_to_update, matching_diligencia))
        lat, lng = get_municipio_coords(process_to_update['municipio'])
        matching_diligencia['lat'] = lat
        matching_diligencia['lng'] = lng

    save_state()
    return jsonify(process_to_update)

@app.route('/api/processos/<int:process_id>', methods=['DELETE'])
def delete_processo(process_id):
    global processos, diligencias
    process_to_delete = next((p for p in processos if p['id'] == process_id), None)
    if process_to_delete is None:
        return jsonify({'error': 'Processo não encontrado'}), 404

    process_numero = process_to_delete['numero']
    processos = [p for p in processos if p['id'] != process_id]
    
    matching_diligencia = next((d for d in diligencias if d.get('process_number') == process_numero), None)
    if matching_diligencia:
        diligencias = [d for d in diligencias if d['id'] != matching_diligencia['id']]

    save_state()
    return jsonify({'success': True}), 200


@app.route('/api/processos/relatorio-docx', methods=['GET'])
def export_processos_docx():
    global report_history
    report_data = sorted(processos, key=lambda item: item.get('numero', ''))
    generated_at = datetime.now().strftime('%d/%m/%Y %H:%M:%S')
    file_stream = build_docx_report(report_data)
    filename = f"relatorio_diligencias_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    report_history.insert(0, {
        'generated_at': generated_at,
        'filename': filename,
        'total_processos': len(report_data),
    })

    save_state()

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/api/relatorios/historico', methods=['GET'])
def get_report_history():
    return jsonify(report_history)

if __name__ == '__main__':
    port = int(os.getenv('PORT', '5000'))
    app.run(host='0.0.0.0', port=port, debug=False)
